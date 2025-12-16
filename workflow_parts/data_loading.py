"""
Data Loading - Workflow Part

Handles PDF extraction (including OCR), DOCX extraction, and initial data loading for RAG workflows.
Supports:
  - Standard PDF files (text extraction)
  - Image-based PDF files (OCR using Tesseract)
  - DOCX files (Word documents)
"""

import fitz
import json
import os
from typing import List, Dict, Union
from pathlib import Path


def extract_text_from_pdf(pdf_path: str, use_ocr: bool = False, auto_ocr: bool = True) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        pdf_path: Path to the PDF file
        use_ocr: If True, use OCR for all pages
        auto_ocr: If True, automatically try OCR if standard extraction yields no text
        
    Returns:
        str: Extracted text from all pages of the PDF
        
    Note:
        For OCR support, install Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki
        Then set TESSERACT_CMD environment variable to the path of tesseract.exe
        Example Windows: C:\Program Files\Tesseract-OCR\tesseract.exe
    """
    pdf_file = fitz.open(pdf_path)
    all_text = ""
    
    for page_num in range(pdf_file.page_count):
        page = pdf_file[page_num]
        text = page.get_text("text")
        
        # If no text found and OCR is enabled (explicit or auto), try OCR
        should_try_ocr = use_ocr or (auto_ocr and not text.strip())
        
        if not text.strip() and should_try_ocr:
            try:
                import pytesseract
                from PIL import Image
                import io
                
                # Configure pytesseract path if environment variable is set
                tesseract_cmd = os.getenv("TESSERACT_CMD")
                if tesseract_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                
                # Render page to image and perform OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution for better OCR
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img, lang='hun+eng')
                if text.strip():
                    print(f"  [Page {page_num + 1}] OCR extracted {len(text)} characters")
            except ImportError:
                print(f"  [Page {page_num + 1}] pytesseract not installed, skipping OCR. Install: pip install pytesseract pillow")
            except Exception as e:
                error_str = str(e)
                if "tesseract is not installed" in error_str.lower():
                    print(f"  [Page {page_num + 1}] Tesseract binary not found. Install from: https://github.com/UB-Mannheim/tesseract/wiki")
                else:
                    print(f"  [Page {page_num + 1}] OCR failed: {e}")
        
        all_text += text
    
    pdf_file.close()
    return all_text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX (Word) file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        str: Extracted text from all paragraphs
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    
    doc = Document(docx_path)
    all_text = ""
    
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
            all_text += "\n"
    
    return all_text


def extract_text_from_file(file_path: str, use_ocr: bool = False, auto_ocr: bool = True) -> str:
    """
    Extract text from any supported file type (PDF, DOCX).
    Automatically detects file type and uses appropriate extraction method.
    
    Args:
        file_path: Path to the file
        use_ocr: For PDFs, enable OCR for all pages
        auto_ocr: For PDFs, automatically try OCR if standard extraction yields no text
        
    Returns:
        str: Extracted text
        
    Raises:
        ValueError: If file type is not supported
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path, use_ocr=use_ocr, auto_ocr=auto_ocr)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Supported: .pdf, .docx")


def load_multiple_files(file_paths: Union[str, List[str]], use_ocr: bool = False, auto_ocr: bool = True, delimiter: str = "\n\n--- Document Boundary ---\n\n") -> Dict[str, str]:
    """
    Load and extract text from multiple files.
    
    Args:
        file_paths: Single file path or list of file paths
        use_ocr: Enable OCR for all PDF pages
        auto_ocr: Automatically try OCR if standard extraction yields no text
        delimiter: String to separate documents in combined text
        
    Returns:
        Dict with keys as filenames and values as extracted text
    """
    if isinstance(file_paths, str):
        file_paths = [file_paths]
    
    results = {}
    for file_path in file_paths:
        try:
            print(f"Loading: {Path(file_path).name}...")
            text = extract_text_from_file(file_path, use_ocr=use_ocr, auto_ocr=auto_ocr)
            results[Path(file_path).name] = text
            print(f"  [OK] Extracted {len(text)} characters")
        except Exception as e:
            print(f"  [ERROR] Error loading {Path(file_path).name}: {e}")
            results[Path(file_path).name] = f"[ERROR: {str(e)}]"
    
    return results


def combine_documents(file_results: Dict[str, str], delimiter: str = "\n\n--- Document Boundary ---\n\n") -> str:
    """
    Combine extracted text from multiple files into a single string.
    
    Args:
        file_results: Dict from load_multiple_files()
        delimiter: String to separate documents
        
    Returns:
        str: Combined text from all files
    """
    combined = []
    for filename, text in file_results.items():
        combined.append(f"[Source: {filename}]\n{text}")
    
    return delimiter.join(combined)



def load_validation_data(val_file: str) -> List[Dict]:
    """
    Load validation queries from a JSON file.
    
    Args:
        val_file: Path to JSON file containing validation queries
        
    Returns:
        List[Dict]: List of query dictionaries with 'question' and 'ideal_answer' keys
    """
    with open(val_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    return data


def extract_queries_from_validation_data(validation_data: List[Dict]) -> List[str]:
    """
    Extract just the questions from validation data.
    
    Args:
        validation_data: List of validation query dicts
        
    Returns:
        List[str]: List of question strings
    """
    return [item['question'] for item in validation_data]
